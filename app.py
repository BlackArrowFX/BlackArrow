import streamlit as st
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# ---------------- 1. INITIALIZE GLOBAL STATE ---------------- #
if "balance" not in st.session_state:
    st.session_state.balance = 2146.11  
if "trades_taken" not in st.session_state:
    st.session_state.trades_taken = 0
if "trade_notes" not in st.session_state:
    st.session_state.trade_notes = ""

# ---------------- SETUP ---------------- #
st.set_page_config(page_title="BlackArrowFX Precision Engine", layout="wide")
now = datetime.now()
dt_string = now.strftime("%d/%m/%Y %H:%M:%S")

# ---------------- WORD REPORT GENERATOR FUNCTION ---------------- #
def create_word_report(data):
    doc = Document()
    
    # Title
    title = doc.add_heading('BlackArrowFX Precision Trade Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Header Info
    doc.add_paragraph(f"Generated on: {dt_string}")
    doc.add_paragraph(f"Instrument: {data['symbol']} | Direction: {data['direction']}")
    
    # Quad Timeframe Table
    doc.add_heading('1. Quad Timeframe Analysis', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'TF'
    hdr_cells[1].text = 'Bias'
    hdr_cells[2].text = 'Swing High'
    hdr_cells[3].text = 'Swing Low'
    
    tf_data = [
        ("4H", data['4h_b'], data['s4h'], data['s4l']),
        ("1H", data['1h_b'], data['s1h'], data['s1l']),
        ("30M", data['30m_b'], data['s30h'], data['s30l']),
        ("15M", data['15m_b'], data['s15h'], data['s15l'])
    ]
    
    for tf, bias, high, low in tf_data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(tf)
        row_cells[1].text = str(bias)
        row_cells[2].text = str(high)
        row_cells[3].text = str(low)

    # Execution Details
    doc.add_heading('2. Execution & Risk', level=1)
    doc.add_paragraph(f"Entry Price: {data['entry']}")
    doc.add_paragraph(f"Stop Loss: {data['sl']}")
    doc.add_paragraph(f"TP1 (1:2): {data['tp1']}")
    
    # Strategy Notes
    doc.add_heading('3. Strategy Notes (Post-Shock Plan)', level=1)
    doc.add_paragraph(data['notes'])
    
    # Save to buffer
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ---------------- SIDEBAR: RISK & SYSTEM ---------------- #
with st.sidebar:
    st.header("⚙️ System Config")
    asset_type = st.selectbox("Select Asset Class", ["METAL (Gold/Silver)", "FOREX", "INDICES / CRYPTO"])
    symbol = st.text_input("Enter Instrument", value="XAUUSD").upper()
    
    st.markdown("---")
    st.header("💰 Risk Engine")
    st.session_state.balance = st.number_input("Current Balance ($)", value=float(st.session_state.balance), format="%.2f")
    
    risk_method = st.radio("Risk Method", ["Percentage (%)", "Fixed Amount ($)"])
    if risk_method == "Percentage (%)":
        risk_pct = st.slider("Risk per Trade (%)", 0.25, 10.0, 1.0)
        current_risk_usd = st.session_state.balance * (risk_pct / 100)
    else:
        current_risk_usd = st.number_input("Risk Amount ($)", min_value=1.0, value=50.0)

    st.markdown("---")
    st.header("🌍 News Filter")
    news_ok = st.toggle("No High Impact News Active", value=False) 

# ---------------- MAIN INTERFACE ---------------- #
st.title(f"🏹 BlackArrowFX: {symbol}")
st.markdown("---")

# ---------------- QUAD TIMEFRAME ANALYSIS ---------------- #
c4h, c1h, c30m, c15m = st.columns(4)
with c4h:
    st.subheader("⏳ 4H BIAS")
    htf_bias = st.radio("4H Trend", ["Bullish ⬆️", "Bearish ⬇️", "Ranging"], key="4h_t")
    s4_h = st.number_input("4H High", value=0.0, format="%.2f")
    s4_l = st.number_input("4H Low", value=0.0, format="%.2f")
with c1h:
    st.subheader("⏱️ 1H STRUC")
    itf_trend = st.radio("1H Trend", ["Bullish ⬆️", "Bearish ⬇️", "Ranging"], key="1h_t")
    s1_h = st.number_input("1H High", value=0.0, format="%.2f")
    s1_l = st.number_input("1H Low", value=0.0, format="%.2f")
with c30m:
    st.subheader("⚡ 30M SHIFT")
    t30_trend = st.radio("30M Trend", ["Bullish ⬆️", "Bearish ⬇️", "Ranging"], key="30m_t")
    s30_h = st.number_input("30M High", value=0.0, format="%.2f")
    s30_l = st.number_input("30M Low", value=0.0, format="%.2f")
with c15m:
    st.subheader("🎯 15M ENTRY")
    t15_trend = st.radio("15M Trend", ["Bullish ⬆️", "Bearish ⬇️", "Ranging"], key="15m_t")
    s15_h = st.number_input("15M High", value=0.0, format="%.2f")
    s15_l = st.number_input("15M Low", value=0.0, format="%.2f")

# ---------------- STRATEGY BOX ---------------- #
st.markdown("---")
st.subheader("📝 POST-SHOCK EXECUTION PLAN")
st.session_state.trade_notes = st.text_area("Strategic Setup:", value=st.session_state.trade_notes, height=400)

# ---------------- PHASE 3: EXECUTE ---------------- #
st.markdown("---")
st.header("🚀 EXECUTION & REPORTING")
col_e, col_r = st.columns(2)

with col_e:
    trade_dir = st.radio("Direction", ["LONG 🔵", "SHORT 🔴"], horizontal=True)
    entry_val = st.number_input("Manual Entry Price", value=0.0, format="%.2f")
    sl_val = st.number_input("Stop Loss", value=0.0, format="%.2f")
    
    pip_factor = 0.1 if asset_type == "METAL (Gold/Silver)" else (0.0001 if asset_type == "FOREX" else 1.0)
    if entry_val > 0 and sl_val > 0:
        actual_pips = abs(entry_val - sl_val) / pip_factor
        tp1 = entry_val + (actual_pips * 2 * pip_factor) if trade_dir == "LONG 🔵" else entry_val - (actual_pips * 2 * pip_factor)
        st.write(f"**TP1 (1:2): {round(tp1, 2)}**")
        st.info(f"🛡️ Management: Set BE and Partial at {round(tp1, 2)}")

with col_r:
    st.subheader("📄 Generate Word Report")
    
    report_data = {
        "symbol": symbol, "direction": trade_dir, "entry": entry_val, "sl": sl_val, "tp1": tp1 if 'tp1' in locals() else 0,
        "4h_b": htf_bias, "s4h": s4_h, "s4l": s4_l,
        "1h_b": itf_trend, "s1h": s1_h, "s1l": s1_l,
        "30m_b": t30_trend, "s30h": s30_h, "s30l": s30_l,
        "15m_b": t15_trend, "s15h": s15_h, "s15l": s15_l,
        "notes": st.session_state.trade_notes
    }

    word_file = create_word_report(report_data)
    
    st.download_button(
        label="📥 DOWNLOAD WORD REPORT (.DOCX)",
        data=word_file,
        file_name=f"Trade_Report_{symbol}_{datetime.now().strftime('%Y%m%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
